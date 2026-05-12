from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "competition"
OUT.mkdir(parents=True, exist_ok=True)


REFERENCES = [
    "Centers for Disease Control and Prevention, Behavioral Risk Factor Surveillance System.",
    "UCI Machine Learning Repository, Diabetes Data Set.",
    "American Diabetes Association, Standards of Care in Diabetes.",
    "World Health Organization, Diabetes fact sheet and digital health guidance.",
]


def add_footnote_marker(paragraph, number: int) -> None:
    run = paragraph.add_run(str(number))
    run.font.superscript = True
    run.font.size = Pt(8)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1
    styles["Normal"].paragraph_format.space_after = Pt(0)
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size in [("Title", 16), ("Heading 1", 13), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_after = Pt(0)


def add_reference_box(document: Document, references: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("Fusnote / reference: ")
    run.bold = True
    for index, reference in enumerate(references, start=1):
        item = paragraph.add_run(f"{index}. {reference} ")
        item.font.size = Pt(9)


def build_agent_description() -> None:
    document = Document()
    style_document(document)
    document.add_heading("Opis agenta Glyco", 0)
    p = document.add_paragraph(
        "Glyco je personalizovani zdravstveni agent za podrsku osobama koje prate rizik i tok dijabetesa tipa 2. "
        "Njegova jedinstvenost je u tome sto nije samo chatbot: agent objedinjuje ML procjenu rizika, ML procjenu monitoring trenda, sigurnosna pravila, smjernice, proaktivne alerte, izvjestaje za doktora i pojednostavljeni porodicni prikaz."
    )
    add_footnote_marker(p, 1)
    document.add_paragraph(
        "Primjena agenta je prakticna: korisnik unosi profil, glukozu, pritisak, tezinu i aktivnost; Glyco zatim objasnjava sta se promijenilo, zasto je vazno, sta uraditi ove sedmice i sta pitati doktora. Sistem ne postavlja dijagnozu i ne zamjenjuje ljekara, nego pomaze u pripremi za razgovor sa zdravstvenim radnikom."
    )
    p = document.add_paragraph(
        "Glyco uci na dva nivoa. Prvi nivo je offline masinsko ucenje: model rizika je treniran na CDC BRFSS datasetu, a model monitoring trenda na UCI diabetes time-series arhivi. Drugi nivo je online adaptacija agenta: korisnik ocjenjuje odgovor, bira preferirani ton i potvrdjuje akciju; agent taj feedback sprema u memoriju i koristi u narednim odgovorima."
    )
    add_footnote_marker(p, 2)
    add_footnote_marker(p, 3)
    p = document.add_paragraph(
        "Takmicarski demo pokazuje kompletan agentski loop: profil korisnika, risk assessment, unos novog loga, promjenu monitoring trenda, agent odgovor, feedback korisnika, personalizovan naredni odgovor i generisanje izvjestaja za doktora ili porodicu."
    )
    add_footnote_marker(p, 4)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Kriterij")
    set_cell_text(table.rows[0].cells[1], "Kako Glyco odgovara")
    rows = [
        ("Jedinstvenost", "Agent spaja ML, pravila, smjernice, memoriju, alerte i family support."),
        ("Primjenjivost", "Korisnik dobija konkretne sedmicne korake i pitanja za doktora."),
        ("Ucenje", "Offline modeli uce iz dataseta; online memorija uci iz feedbacka korisnika."),
        ("Kod", "Backend, frontend, ML pipeline, testovi i dokumentacija su odvojeni i pokriveni."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)
    add_reference_box(document, REFERENCES)
    document.save(OUT / "Opis-agenta-Glyco.docx")


def build_code_documentation() -> None:
    document = Document()
    style_document(document)
    document.add_heading("Dokumentacija koda za Glyco", 0)
    sections = [
        ("1. Arhitektura sistema", "Glyco je podijeljen u React frontend, FastAPI backend, SQLite bazu i ML pipeline. Frontend prikazuje dashboard, unos profila, monitoring logove, izvjestaje, family view i agent chat. Backend prima podatke, cuva ih u bazi, poziva modele, generise objasnjenja i vraca strukturirane API odgovore. SQLite cuva korisnike, profile, logove, procjene, izvjestaje, alerte i feedback agenta."),
        ("2. Backend i agent", "Agent je implementiran u backend/app/agent. Modul tools.py ucitava profil, logove, risk assessment, trend assessment i smjernice. Modul safety.py provjerava urgentne simptome prije generisanja odgovora. Modul agent_service.py orkestrira cijeli tok: kontekst korisnika, memoriju, modele, smjernice, sigurnosnu granicu i LLM/fallback odgovor."),
        ("3. Learning loop", "Tabela agent_feedback cuva ocjenu korisnosti, preferirani ton, potvrdjenu akciju i napomenu. Agent cita zadnjih 12 feedback zapisa i izracunava learning_summary: broj signala, helpful rate, preferirani ton i potvrdjene akcije. Taj summary se vraca frontend-u i koristi u narednom odgovoru, cime se dokazuje online adaptacija."),
        ("4. Masinsko ucenje", "Risk model koristi BRFSS dataset i random forest klasifikator. Profil se prevodi u feature red koji model ocekuje, a rezultat se pretvara u low, medium ili high rizik. Monitoring model koristi UCI time-series podatke i dnevne karakteristike kao sto su prosjecna glukoza, varijabilnost, high/low count, aktivnost i kratkorocni slope."),
        ("5. Fallback i sigurnost", "Ako ML artefakti nisu dostupni ili korisnik nema dovoljno historije, backend prelazi na deterministic rules fallback. Sistem ostaje funkcionalan i transparentno pokazuje koji model ili fallback je koristen. Urgentni simptomi zaobilaze normalan coaching i vracaju poruku za hitno obracanje ljekaru ili hitnoj sluzbi."),
        ("6. Frontend interface", "Agent ekran prikazuje chat, brza pitanja, tool calls, guideline grounding, safety note i Agent Memory karticu. Teach Glyco panel omogucava korisniku da sacuva feedback i time personalizuje naredni odgovor. Ostale stranice pokazuju risk, monitoring, reports, care plan i family support."),
        ("7. Testiranje", "Backend testovi provjeravaju ML artefakte, demo korisnike, monitoring model, fallback za nedovoljnu historiju, izvjestaje, agent tool calls, urgent safety, proactive alerts i novi feedback/memory tok. Time su pokriveni najvazniji kriteriji: primjenjivost, kod, dokumentacija i kompleksnost ucenja."),
    ]
    for title, body in sections:
        document.add_heading(title, level=1)
        p = document.add_paragraph(body)
        if title in {"3. Learning loop", "4. Masinsko ucenje"}:
            add_footnote_marker(p, 1 if "BRFSS" in body else 2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Sloj", "Glavni fajlovi", "Uloga"]):
        set_cell_text(cell, text)
    rows = [
        ("API", "backend/app/api/routes.py", "Endpointi za profile, logove, procjene, izvjestaje, agent i feedback."),
        ("Agent", "backend/app/agent", "Orkestracija alata, memorija, sigurnost i odgovor korisniku."),
        ("ML", "backend/app/ml/inference.py; ml/scripts", "Ucitavanje artefakata, feature engineering, trening i fallback."),
        ("Frontend", "frontend/src/pages/Agent.tsx", "Chat interface, tool evidence, memory kartica i Teach Glyco feedback."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    add_reference_box(document, REFERENCES)
    document.save(OUT / "Dokumentacija-koda-Glyco.docx")


if __name__ == "__main__":
    build_agent_description()
    build_code_documentation()
